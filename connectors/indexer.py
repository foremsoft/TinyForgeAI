"""
RAG Document Indexer for TinyForgeAI.

This module provides document indexing capabilities for Retrieval-Augmented Generation (RAG).
It supports chunking documents, generating embeddings, and storing in vector databases.

Usage:
    from connectors.indexer import DocumentIndexer, IndexerConfig

    config = IndexerConfig(
        chunk_size=512,
        chunk_overlap=50,
        embedding_model="sentence-transformers/all-MiniLM-L6-v2"
    )
    indexer = DocumentIndexer(config)
    indexer.index_file("document.pdf")
    results = indexer.search("What is machine learning?", top_k=5)
"""

import hashlib
import json
import logging
import os
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union

logger = logging.getLogger(__name__)

# Check for optional dependencies
EMBEDDINGS_AVAILABLE = False
try:
    from sentence_transformers import SentenceTransformer
    EMBEDDINGS_AVAILABLE = True
except ImportError:
    logger.warning(
        "sentence-transformers not installed. "
        "Install with: pip install sentence-transformers"
    )

NUMPY_AVAILABLE = False
try:
    import numpy as np
    NUMPY_AVAILABLE = True
except ImportError:
    logger.warning("numpy not installed. Install with: pip install numpy")


@dataclass
class IndexerConfig:
    """Configuration for document indexing."""

    # Chunking settings
    chunk_size: int = 512  # Characters per chunk
    chunk_overlap: int = 50  # Overlap between chunks

    # Embedding settings
    embedding_model: str = "sentence-transformers/all-MiniLM-L6-v2"
    embedding_dimension: int = 384  # Dimension for MiniLM

    # Storage settings
    index_path: str = "./vector_index"
    use_faiss: bool = False  # Use FAISS for vector search

    # Processing settings
    batch_size: int = 32
    normalize_embeddings: bool = True


@dataclass
class Document:
    """A document chunk with metadata."""
    id: str
    content: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    embedding: Optional[List[float]] = None


@dataclass
class SearchResult:
    """Search result with score."""
    document: Document
    score: float


class TextChunker:
    """Split text into overlapping chunks."""

    def __init__(self, chunk_size: int = 512, overlap: int = 50):
        self.chunk_size = chunk_size
        self.overlap = overlap

    def chunk(self, text: str) -> List[str]:
        """
        Split text into overlapping chunks.

        Args:
            text: Input text to chunk.

        Returns:
            List of text chunks.
        """
        if not text:
            return []

        chunks = []
        start = 0

        while start < len(text):
            end = start + self.chunk_size

            # Try to break at sentence boundary
            if end < len(text):
                # Look for sentence endings
                for sep in ['. ', '.\n', '! ', '? ', '\n\n']:
                    last_sep = text[start:end].rfind(sep)
                    if last_sep > self.chunk_size // 2:
                        end = start + last_sep + len(sep)
                        break

            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)

            # Move start with overlap
            start = end - self.overlap

        return chunks

    def chunk_with_metadata(
        self,
        text: str,
        source: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> List[Document]:
        """
        Chunk text and create Document objects with metadata.

        Args:
            text: Input text.
            source: Source identifier (e.g., filename).
            metadata: Additional metadata.

        Returns:
            List of Document objects.
        """
        chunks = self.chunk(text)
        documents = []

        for i, chunk in enumerate(chunks):
            doc_id = hashlib.md5(f"{source}:{i}:{chunk[:50]}".encode()).hexdigest()[:12]

            doc_metadata = {
                "source": source,
                "chunk_index": i,
                "total_chunks": len(chunks),
                "created_at": datetime.utcnow().isoformat(),
            }
            if metadata:
                doc_metadata.update(metadata)

            documents.append(Document(
                id=doc_id,
                content=chunk,
                metadata=doc_metadata,
            ))

        return documents


class DocumentIndexer:
    """
    Index documents for RAG retrieval.

    Supports:
    - Text chunking with overlap
    - Embedding generation (sentence-transformers)
    - Simple in-memory vector search
    - FAISS integration (optional)
    """

    def __init__(self, config: Optional[IndexerConfig] = None):
        """
        Initialize the indexer.

        Args:
            config: Indexer configuration.
        """
        self.config = config or IndexerConfig()
        self.chunker = TextChunker(
            chunk_size=self.config.chunk_size,
            overlap=self.config.chunk_overlap
        )
        self.documents: Dict[str, Document] = {}
        self.embeddings: Optional[Any] = None  # numpy array
        self.embedding_model = None

        # Create index directory
        Path(self.config.index_path).mkdir(parents=True, exist_ok=True)

    def _load_embedding_model(self):
        """Load the embedding model lazily."""
        if self.embedding_model is not None:
            return

        if not EMBEDDINGS_AVAILABLE:
            raise RuntimeError(
                "sentence-transformers not installed. "
                "Install with: pip install sentence-transformers"
            )

        logger.info(f"Loading embedding model: {self.config.embedding_model}")
        self.embedding_model = SentenceTransformer(self.config.embedding_model)

    def _compute_embeddings(self, texts: List[str]) -> Any:
        """
        Compute embeddings for a list of texts.

        Args:
            texts: List of text strings.

        Returns:
            Numpy array of embeddings.
        """
        self._load_embedding_model()

        embeddings = self.embedding_model.encode(
            texts,
            batch_size=self.config.batch_size,
            normalize_embeddings=self.config.normalize_embeddings,
            show_progress_bar=len(texts) > 100,
        )

        return embeddings

    def index_text(
        self,
        text: str,
        source: str = "text",
        metadata: Optional[Dict[str, Any]] = None
    ) -> List[str]:
        """
        Index a text string.

        Args:
            text: Text to index.
            source: Source identifier.
            metadata: Additional metadata.

        Returns:
            List of document IDs that were indexed.
        """
        documents = self.chunker.chunk_with_metadata(text, source, metadata)

        if not documents:
            return []

        # Compute embeddings
        texts = [doc.content for doc in documents]
        embeddings = self._compute_embeddings(texts)

        # Store documents with embeddings
        doc_ids = []
        for doc, embedding in zip(documents, embeddings):
            doc.embedding = embedding.tolist()
            self.documents[doc.id] = doc
            doc_ids.append(doc.id)

        # Update index
        self._rebuild_index()

        logger.info(f"Indexed {len(doc_ids)} chunks from '{source}'")
        return doc_ids

    def index_file(
        self,
        file_path: Union[str, Path],
        metadata: Optional[Dict[str, Any]] = None
    ) -> List[str]:
        """
        Index a file (supports txt, pdf, docx, json, jsonl).

        Args:
            file_path: Path to file.
            metadata: Additional metadata.

        Returns:
            List of document IDs.
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Get file metadata
        file_metadata = {
            "filename": file_path.name,
            "file_type": file_path.suffix.lower(),
            "file_size": file_path.stat().st_size,
        }
        if metadata:
            file_metadata.update(metadata)

        # Read content based on file type
        suffix = file_path.suffix.lower()

        if suffix in ['.txt', '.md']:
            text = file_path.read_text(encoding='utf-8')
        elif suffix == '.json':
            data = json.loads(file_path.read_text(encoding='utf-8'))
            text = json.dumps(data, indent=2)
        elif suffix == '.jsonl':
            lines = file_path.read_text(encoding='utf-8').strip().split('\n')
            records = [json.loads(line) for line in lines if line.strip()]
            text = '\n\n'.join(
                f"Input: {r.get('input', '')}\nOutput: {r.get('output', '')}"
                for r in records
            )
        elif suffix == '.pdf':
            text = self._extract_pdf(file_path)
        elif suffix == '.docx':
            text = self._extract_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

        return self.index_text(text, source=str(file_path), metadata=file_metadata)

    def _extract_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            text = '\n\n'.join(page.get_text() for page in doc)
            doc.close()
            return text
        except ImportError:
            raise RuntimeError(
                "PyMuPDF not installed. Install with: pip install PyMuPDF"
            )

    def _extract_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            return '\n\n'.join(para.text for para in doc.paragraphs)
        except ImportError:
            raise RuntimeError(
                "python-docx not installed. Install with: pip install python-docx"
            )

    def _rebuild_index(self):
        """Rebuild the embedding index."""
        if not NUMPY_AVAILABLE:
            return

        if not self.documents:
            self.embeddings = None
            return

        # Stack all embeddings
        doc_list = list(self.documents.values())
        self.embeddings = np.array([doc.embedding for doc in doc_list])

    def search(
        self,
        query: str,
        top_k: int = 5,
        threshold: float = 0.0
    ) -> List[SearchResult]:
        """
        Search for similar documents.

        Args:
            query: Search query.
            top_k: Number of results to return.
            threshold: Minimum similarity score.

        Returns:
            List of SearchResult objects.
        """
        if not self.documents:
            return []

        if not NUMPY_AVAILABLE:
            raise RuntimeError("numpy required for search")

        # Compute query embedding
        query_embedding = self._compute_embeddings([query])[0]

        # Compute similarities
        similarities = np.dot(self.embeddings, query_embedding)

        # Get top-k indices
        top_indices = np.argsort(similarities)[::-1][:top_k]

        # Build results
        doc_list = list(self.documents.values())
        results = []

        for idx in top_indices:
            score = float(similarities[idx])
            if score >= threshold:
                results.append(SearchResult(
                    document=doc_list[idx],
                    score=score
                ))

        return results

    def search_with_context(
        self,
        query: str,
        top_k: int = 5,
        context_window: int = 1
    ) -> str:
        """
        Search and return formatted context string for RAG.

        Args:
            query: Search query.
            top_k: Number of results.
            context_window: Number of adjacent chunks to include.

        Returns:
            Formatted context string.
        """
        results = self.search(query, top_k=top_k)

        context_parts = []
        for i, result in enumerate(results):
            context_parts.append(
                f"[{i+1}] (score: {result.score:.3f})\n{result.document.content}"
            )

        return "\n\n---\n\n".join(context_parts)

    def save_index(self, path: Optional[str] = None):
        """
        Save the index to disk.

        Args:
            path: Output path. Uses config.index_path if not specified.
        """
        path = Path(path or self.config.index_path)
        path.mkdir(parents=True, exist_ok=True)

        # Save documents
        docs_data = {
            doc_id: {
                "id": doc.id,
                "content": doc.content,
                "metadata": doc.metadata,
                "embedding": doc.embedding,
            }
            for doc_id, doc in self.documents.items()
        }

        with open(path / "documents.json", "w") as f:
            json.dump(docs_data, f)

        # Save config
        with open(path / "config.json", "w") as f:
            json.dump({
                "chunk_size": self.config.chunk_size,
                "chunk_overlap": self.config.chunk_overlap,
                "embedding_model": self.config.embedding_model,
                "embedding_dimension": self.config.embedding_dimension,
            }, f)

        logger.info(f"Saved index with {len(self.documents)} documents to {path}")

    def load_index(self, path: Optional[str] = None):
        """
        Load the index from disk.

        Args:
            path: Index path. Uses config.index_path if not specified.
        """
        path = Path(path or self.config.index_path)

        # Load documents
        docs_path = path / "documents.json"
        if not docs_path.exists():
            raise FileNotFoundError(f"No index found at {path}")

        with open(docs_path) as f:
            docs_data = json.load(f)

        self.documents = {
            doc_id: Document(
                id=data["id"],
                content=data["content"],
                metadata=data["metadata"],
                embedding=data["embedding"],
            )
            for doc_id, data in docs_data.items()
        }

        # Rebuild numpy index
        self._rebuild_index()

        logger.info(f"Loaded index with {len(self.documents)} documents from {path}")

    def clear(self):
        """Clear all indexed documents."""
        self.documents = {}
        self.embeddings = None
        logger.info("Cleared index")

    def get_stats(self) -> Dict[str, Any]:
        """Get index statistics."""
        return {
            "total_documents": len(self.documents),
            "embedding_dimension": self.config.embedding_dimension,
            "embedding_model": self.config.embedding_model,
            "chunk_size": self.config.chunk_size,
            "index_path": self.config.index_path,
        }


# ============================================
# CLI Interface
# ============================================

def main():
    """CLI entry point for document indexing."""
    import argparse

    parser = argparse.ArgumentParser(
        description="Index documents for RAG retrieval"
    )
    subparsers = parser.add_subparsers(dest="command", required=True)

    # Index command
    index_parser = subparsers.add_parser("index", help="Index a file")
    index_parser.add_argument("file", help="File to index")
    index_parser.add_argument(
        "--index-path", "-o",
        default="./vector_index",
        help="Output index path"
    )
    index_parser.add_argument(
        "--chunk-size",
        type=int,
        default=512,
        help="Chunk size in characters"
    )

    # Search command
    search_parser = subparsers.add_parser("search", help="Search the index")
    search_parser.add_argument("query", help="Search query")
    search_parser.add_argument(
        "--index-path", "-i",
        default="./vector_index",
        help="Index path"
    )
    search_parser.add_argument(
        "--top-k", "-k",
        type=int,
        default=5,
        help="Number of results"
    )

    # Stats command
    stats_parser = subparsers.add_parser("stats", help="Show index stats")
    stats_parser.add_argument(
        "--index-path", "-i",
        default="./vector_index",
        help="Index path"
    )

    args = parser.parse_args()

    # Setup logging
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s - %(levelname)s - %(message)s"
    )

    if args.command == "index":
        config = IndexerConfig(
            chunk_size=args.chunk_size,
            index_path=args.index_path,
        )
        indexer = DocumentIndexer(config)

        # Try to load existing index
        try:
            indexer.load_index()
        except FileNotFoundError:
            pass

        # Index file
        doc_ids = indexer.index_file(args.file)
        print(f"Indexed {len(doc_ids)} chunks from {args.file}")

        # Save
        indexer.save_index()
        print(f"Index saved to {args.index_path}")

    elif args.command == "search":
        config = IndexerConfig(index_path=args.index_path)
        indexer = DocumentIndexer(config)
        indexer.load_index()

        results = indexer.search(args.query, top_k=args.top_k)

        print(f"\nSearch results for: '{args.query}'\n")
        for i, result in enumerate(results):
            print(f"[{i+1}] Score: {result.score:.4f}")
            print(f"    Source: {result.document.metadata.get('source', 'unknown')}")
            print(f"    Content: {result.document.content[:200]}...")
            print()

    elif args.command == "stats":
        config = IndexerConfig(index_path=args.index_path)
        indexer = DocumentIndexer(config)
        indexer.load_index()

        stats = indexer.get_stats()
        print("\nIndex Statistics:")
        for key, value in stats.items():
            print(f"  {key}: {value}")


if __name__ == "__main__":
    main()
